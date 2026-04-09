from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path("/Users/wangzixing/Desktop/各种软件集合/智绘阅读/zhihui-reading")
DOC_DIR = ROOT / "output" / "doc"
TARGET_FILES = [
    "开发文档.docx",
    "测试文档.docx",
    "设计及创新性分析报告.docx",
    "技术研究报告.docx",
]


def set_run_fonts(run, ascii_font: str, east_asia_font: str, size_pt: int | None = None, bold: bool | None = None):
    run.font.name = ascii_font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia_font)
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if bold is not None:
        run.bold = bold


def set_paragraph_shading(paragraph, fill: str):
    p_pr = paragraph._element.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def ensure_style(document: Document, style_name: str):
    try:
      return document.styles[style_name]
    except KeyError:
      return document.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)


def configure_document_styles(document: Document):
    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.first_line_indent = Cm(0.74)

    for style_name in ["Body Text", "First Paragraph", "Compact"]:
        if style_name in document.styles:
            style = document.styles[style_name]
            style.font.name = "Times New Roman"
            style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
            style.font.size = Pt(12)
            style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            style.paragraph_format.space_after = Pt(0)
            style.paragraph_format.first_line_indent = Cm(0.74 if style_name != "Compact" else 0)

    title_style = document.styles["Title"]
    title_style.font.name = "Times New Roman"
    title_style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    title_style.font.size = Pt(20)
    title_style.font.bold = True
    title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_style.paragraph_format.space_after = Pt(12)

    h1 = document.styles["Heading 1"]
    h1.font.name = "Times New Roman"
    h1._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    h1.font.size = Pt(16)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0x11, 0x11, 0x11)
    h1.paragraph_format.space_before = Pt(12)
    h1.paragraph_format.space_after = Pt(6)

    h2 = document.styles["Heading 2"]
    h2.font.name = "Times New Roman"
    h2._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    h2.font.size = Pt(14)
    h2.font.bold = True
    h2.paragraph_format.space_before = Pt(10)
    h2.paragraph_format.space_after = Pt(4)

    h3 = document.styles["Heading 3"]
    h3.font.name = "Times New Roman"
    h3._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    h3.font.size = Pt(12)
    h3.font.bold = True
    h3.paragraph_format.space_before = Pt(8)
    h3.paragraph_format.space_after = Pt(4)

    if "Image Caption" in document.styles:
        cap = document.styles["Image Caption"]
        cap.font.name = "Times New Roman"
        cap._element.rPr.rFonts.set(qn("w:eastAsia"), "楷体")
        cap.font.size = Pt(10.5)
        cap.font.italic = True
        cap.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_after = Pt(6)
        cap.paragraph_format.first_line_indent = Cm(0)

    if "Captioned Figure" in document.styles:
        fig = document.styles["Captioned Figure"]
        fig.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fig.paragraph_format.first_line_indent = Cm(0)
        fig.paragraph_format.space_after = Pt(0)

    if "Source Code" in document.styles:
        code = document.styles["Source Code"]
        code.font.name = "Courier New"
        code._element.rPr.rFonts.set(qn("w:eastAsia"), "等宽更纱黑体 SC")
        code.font.size = Pt(9)
        code.paragraph_format.left_indent = Cm(0.6)
        code.paragraph_format.right_indent = Cm(0.6)
        code.paragraph_format.space_before = Pt(4)
        code.paragraph_format.space_after = Pt(4)
        code.paragraph_format.first_line_indent = Cm(0)


def configure_sections(document: Document):
    for section in document.sections:
        section.start_type = WD_SECTION.NEW_PAGE if section is not document.sections[0] else section.start_type
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.5)
        section.header_distance = Cm(1.5)
        section.footer_distance = Cm(1.5)


def style_paragraphs(document: Document):
    if document.paragraphs:
        first = document.paragraphs[0]
        first.style = document.styles["Title"]
        first.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in first.runs:
            set_run_fonts(run, "Times New Roman", "黑体", 20, True)

        if len(document.paragraphs) > 1 and "版本：" in document.paragraphs[1].text:
            meta = document.paragraphs[1]
            meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
            meta.paragraph_format.space_after = Pt(10)
            meta.paragraph_format.first_line_indent = Cm(0)
            for run in meta.runs:
                set_run_fonts(run, "Times New Roman", "宋体", 10)
                run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    for paragraph in document.paragraphs:
        style_name = paragraph.style.name
        text = paragraph.text.strip()

        if style_name in {"Normal", "Body Text", "First Paragraph", "Compact"}:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            if style_name == "Compact":
                paragraph.paragraph_format.first_line_indent = Cm(0)
                paragraph.paragraph_format.left_indent = Cm(0.74)
            else:
                paragraph.paragraph_format.first_line_indent = Cm(0.74)
            for run in paragraph.runs:
                set_run_fonts(run, "Times New Roman", "宋体", 12)

        if style_name == "Source Code":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.first_line_indent = Cm(0)
            set_paragraph_shading(paragraph, "F3F4F6")
            for run in paragraph.runs:
                set_run_fonts(run, "Courier New", "等宽更纱黑体 SC", 9)

        if style_name == "Image Caption":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.first_line_indent = Cm(0)
            for run in paragraph.runs:
                set_run_fonts(run, "Times New Roman", "楷体", 10)
                run.font.italic = True

        if style_name == "Captioned Figure":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.first_line_indent = Cm(0)

        if not text:
            paragraph.paragraph_format.space_after = Pt(0)


def style_tables(document: Document):
    for table in document.tables:
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for row_idx, row in enumerate(table.rows):
            for cell in row.cells:
                if row_idx == 0:
                    set_cell_shading(cell, "D9EAF7")
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if row_idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
                    paragraph.paragraph_format.first_line_indent = Cm(0)
                    for run in paragraph.runs:
                        set_run_fonts(run, "Times New Roman", "宋体", 10 if row_idx == 0 else 10)
                        if row_idx == 0:
                            run.bold = True


def process_docx(doc_path: Path):
    document = Document(doc_path)
    configure_sections(document)
    configure_document_styles(document)
    style_paragraphs(document)
    style_tables(document)
    document.save(doc_path)


def main():
    for file_name in TARGET_FILES:
        process_docx(DOC_DIR / file_name)
        print(f"Styled {file_name}")


if __name__ == "__main__":
    main()
